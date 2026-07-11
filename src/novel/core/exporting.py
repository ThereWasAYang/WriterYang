from __future__ import annotations

from dataclasses import dataclass
from datetime import datetime
import hashlib
import os
from pathlib import Path
import re
import tempfile
from typing import Literal

from docx import Document

from novel.core.drafting import _chapter_number_text
from novel.core.io import atomic_write_model_json, atomic_write_text, backup_if_exists, load_json_model, load_yaml_model
from novel.core.lifecycle import LifecycleError, accepted_chapter_commit
from novel.core.polishing import DraftDocument, PolishingError, read_markdown_with_front_matter
from novel.core.schemas import ExportManifest, ExportRecord, ExportSourceChapter, ProjectConfig
from novel.core.timeutil import utc_now_precise

ExportType = Literal["markdown", "docx", "html", "txt"]


class ExportError(RuntimeError):
    """Raised when export cannot proceed safely."""


@dataclass(frozen=True)
class MarkdownExportOptions:
    root: Path
    chapters: tuple[int, ...] = ()
    from_chapter: int | None = None
    to_chapter: int | None = None
    output_path: Path | None = None
    title: str | None = None
    include_toc: bool = False
    volume_title: str | None = None
    chapter_number_style: Literal["chinese", "arabic", "chapter", "plain"] = "chinese"
    force: bool = False


@dataclass(frozen=True)
class DocxExportOptions:
    root: Path
    chapters: tuple[int, ...] = ()
    from_chapter: int | None = None
    to_chapter: int | None = None
    output_path: Path | None = None
    title: str | None = None
    force: bool = False


@dataclass(frozen=True)
class ExportedChapter:
    chapter_number: int
    title: str
    path: Path
    document: DraftDocument
    acceptance_commit_id: str
    candidate_artifact_id: str
    audit_artifact_id: str
    state_proposal_artifact_id: str
    post_state_sha256: str
    post_timeline_sha256: str


@dataclass(frozen=True)
class MarkdownExportResult:
    output_path: Path
    manifest_path: Path
    manifest: ExportManifest
    exported_chapters: tuple[int, ...]
    warnings: tuple[str, ...] = ()


@dataclass(frozen=True)
class DocxExportResult:
    output_path: Path
    manifest_path: Path
    manifest: ExportManifest
    exported_chapters: tuple[int, ...]
    warnings: tuple[str, ...] = ()


def export_markdown(options: MarkdownExportOptions) -> MarkdownExportResult:
    root = options.root.resolve()
    output_path = _resolve_output_path(root, options.output_path, default_name="novel.md")
    manifest_path = root / "exports" / "export_manifest.json"
    if output_path.exists() and not options.force:
        raise ExportError(f"{output_path} already exists; use --force to overwrite it")

    project = load_yaml_model(root / "project.yaml", ProjectConfig)
    title = options.title.strip() if options.title and options.title.strip() else project.title
    chapters, warnings = collect_export_chapters(options, root)
    if not chapters:
        raise ExportError("no chapters selected for export")

    output_path.parent.mkdir(parents=True, exist_ok=True)
    if options.force:
        backup_if_exists(output_path, reason="force")
    atomic_write_text(
        output_path,
        render_markdown_export(
            title,
            chapters,
            include_toc=options.include_toc,
            volume_title=options.volume_title,
            chapter_number_style=options.chapter_number_style,
        ),
    )
    manifest = update_export_manifest(
        root=root,
        manifest_path=manifest_path,
        output_path=output_path,
        export_type="markdown",
        title=title,
        chapters=chapters,
    )
    return MarkdownExportResult(
        output_path=output_path,
        manifest_path=manifest_path,
        manifest=manifest,
        exported_chapters=tuple(chapter.chapter_number for chapter in chapters),
        warnings=tuple(warnings),
    )


def export_docx(options: DocxExportOptions) -> DocxExportResult:
    root = options.root.resolve()
    output_path = _resolve_output_path(root, options.output_path, default_name="novel.docx")
    manifest_path = root / "exports" / "export_manifest.json"
    if output_path.exists() and not options.force:
        raise ExportError(f"{output_path} already exists; use --force to overwrite it")

    project = load_yaml_model(root / "project.yaml", ProjectConfig)
    title = options.title.strip() if options.title and options.title.strip() else project.title
    chapters, warnings = collect_export_chapters(_docx_to_markdown_options(options), root)
    if not chapters:
        raise ExportError("no chapters selected for export")

    output_path.parent.mkdir(parents=True, exist_ok=True)
    if options.force:
        backup_if_exists(output_path, reason="force")
    render_docx_export(title, chapters, output_path)
    manifest = update_export_manifest(
        root=root,
        manifest_path=manifest_path,
        output_path=output_path,
        export_type="docx",
        title=title,
        chapters=chapters,
    )
    return DocxExportResult(
        output_path=output_path,
        manifest_path=manifest_path,
        manifest=manifest,
        exported_chapters=tuple(chapter.chapter_number for chapter in chapters),
        warnings=tuple(warnings),
    )


def collect_export_chapters(options: MarkdownExportOptions, root: Path) -> tuple[list[ExportedChapter], list[str]]:
    selected_numbers = _selected_chapter_numbers(options, root)
    warnings: list[str] = []
    chapters: list[ExportedChapter] = []
    for chapter_number in selected_numbers:
        chapter_dir = root / "memory" / "chapters" / f"{chapter_number:03d}"
        path = chapter_dir / "accepted.md"
        if not path.exists():
            warnings.append(f"chapter {chapter_number} has no accepted.md; skipped")
            continue
        try:
            acceptance = accepted_chapter_commit(root, chapter_number)
        except LifecycleError as exc:
            raise ExportError(str(exc)) from exc
        document = _read_chapter_document(path)
        metadata_number = document.metadata.get("chapter_number")
        if metadata_number != chapter_number:
            raise ExportError(
                f"{path} chapter_number {metadata_number} does not match directory chapter {chapter_number}"
            )
        chapters.append(
            ExportedChapter(
                chapter_number=chapter_number,
                title=str(document.metadata.get("title") or f"Chapter {chapter_number}"),
                path=path,
                document=document,
                acceptance_commit_id=acceptance.commit_id,
                candidate_artifact_id=acceptance.candidate.artifact_id,
                audit_artifact_id=acceptance.audit.artifact_id,
                state_proposal_artifact_id=acceptance.state_proposal.artifact_id,
                post_state_sha256=acceptance.post_state_sha256,
                post_timeline_sha256=acceptance.post_timeline_sha256,
            )
        )
    return chapters, warnings


def render_markdown_export(
    title: str,
    chapters: list[ExportedChapter],
    *,
    include_toc: bool = False,
    volume_title: str | None = None,
    chapter_number_style: Literal["chinese", "arabic", "chapter", "plain"] = "chinese",
) -> str:
    lines = [f"# {title}", ""]
    if include_toc:
        lines.extend(["## 目录", ""])
        if volume_title and volume_title.strip():
            lines.extend([f"- {volume_title.strip()}"])
            for chapter in chapters:
                heading = _chapter_heading(chapter, chapter_number_style)
                lines.append(f"  - [{heading}](#{_markdown_anchor(heading)})")
        else:
            for chapter in chapters:
                heading = _chapter_heading(chapter, chapter_number_style)
                lines.append(f"- [{heading}](#{_markdown_anchor(heading)})")
        lines.append("")
    if volume_title and volume_title.strip():
        lines.extend([f"## {volume_title.strip()}", ""])
    for chapter in chapters:
        heading = _chapter_heading(chapter, chapter_number_style)
        lines.extend(
            [
                f"## {heading}",
                "",
                _strip_leading_heading(chapter.document.body),
                "",
            ]
        )
    return "\n".join(lines).rstrip() + "\n"


def render_docx_export(title: str, chapters: list[ExportedChapter], output_path: Path) -> None:
    document = Document()
    document.add_heading(title, level=0)
    for chapter in chapters:
        document.add_heading(
            f"第{_chapter_number_text(chapter.chapter_number)}章 {chapter.title}",
            level=1,
        )
        body = _strip_leading_heading(chapter.document.body)
        for line in body.splitlines():
            if not line.strip():
                document.add_paragraph("")
                continue
            if line.startswith("### "):
                document.add_heading(line[4:].strip(), level=3)
            elif line.startswith("## "):
                document.add_heading(line[3:].strip(), level=2)
            elif line.startswith("# "):
                document.add_heading(line[2:].strip(), level=1)
            else:
                document.add_paragraph(line)
    fd, temp_name = tempfile.mkstemp(prefix=f".{output_path.name}.", suffix=".tmp", dir=str(output_path.parent))
    os.close(fd)
    try:
        document.save(temp_name)
        os.replace(temp_name, output_path)
    finally:
        Path(temp_name).unlink(missing_ok=True)


def update_export_manifest(
    *,
    root: Path,
    manifest_path: Path,
    output_path: Path,
    export_type: ExportType,
    title: str,
    chapters: list[ExportedChapter],
) -> ExportManifest:
    if manifest_path.exists():
        manifest = load_json_model(manifest_path, ExportManifest)
    else:
        manifest = ExportManifest()
    record = ExportRecord(
        id=_export_id(utc_now_precise()),
        type=export_type,
        source_chapters=[chapter.chapter_number for chapter in chapters],
        source_chapter_details=[
            ExportSourceChapter(
                chapter_number=chapter.chapter_number,
                title=chapter.title,
                path=_rel(root, chapter.path),
                accepted=True,
                sha256=_sha256_file(chapter.path),
                acceptance_commit_id=chapter.acceptance_commit_id,
                candidate_artifact_id=chapter.candidate_artifact_id,
                audit_artifact_id=chapter.audit_artifact_id,
                state_proposal_artifact_id=chapter.state_proposal_artifact_id,
                post_state_sha256=chapter.post_state_sha256,
                post_timeline_sha256=chapter.post_timeline_sha256,
            )
            for chapter in chapters
        ],
        output_path=_rel(root, output_path),
        created_at=utc_now_precise(),
        title=title,
    )
    manifest.exports.append(record)
    manifest_path.parent.mkdir(parents=True, exist_ok=True)
    backup_if_exists(manifest_path, reason="export_manifest")
    atomic_write_model_json(manifest_path, manifest)
    return manifest


def parse_chapter_selector(chapters: str | None) -> tuple[int, ...]:
    if not chapters:
        return ()
    values: list[int] = []
    for raw in chapters.split(","):
        item = raw.strip()
        if not item:
            continue
        try:
            value = int(item)
        except ValueError as exc:
            raise ExportError(f"invalid chapter number: {item}") from exc
        if value < 1:
            raise ExportError("chapter numbers must be positive integers")
        values.append(value)
    return tuple(values)


def _docx_to_markdown_options(options: DocxExportOptions) -> MarkdownExportOptions:
    return MarkdownExportOptions(
        root=options.root,
        chapters=options.chapters,
        from_chapter=options.from_chapter,
        to_chapter=options.to_chapter,
        output_path=options.output_path,
        title=options.title,
        force=options.force,
    )


def _selected_chapter_numbers(options: MarkdownExportOptions, root: Path) -> list[int]:
    if options.chapters:
        return sorted(set(options.chapters))
    chapters_dir = root / "memory" / "chapters"
    numbers: list[int] = []
    if chapters_dir.exists():
        for child in chapters_dir.iterdir():
            if child.is_dir() and child.name.isdigit():
                numbers.append(int(child.name))
    if options.from_chapter is not None:
        if options.from_chapter < 1:
            raise ExportError("--from must be a positive integer")
        numbers = [number for number in numbers if number >= options.from_chapter]
    if options.to_chapter is not None:
        if options.to_chapter < 1:
            raise ExportError("--to must be a positive integer")
        numbers = [number for number in numbers if number <= options.to_chapter]
    if (
        options.from_chapter is not None
        and options.to_chapter is not None
        and options.from_chapter > options.to_chapter
    ):
        raise ExportError("--from must be less than or equal to --to")
    return sorted(set(numbers))


def _read_chapter_document(path: Path) -> DraftDocument:
    try:
        return read_markdown_with_front_matter(path)
    except PolishingError as exc:
        raise ExportError(str(exc)) from exc


def _strip_leading_heading(body: str) -> str:
    lines = body.strip().splitlines()
    while lines and not lines[0].strip():
        lines.pop(0)
    if lines and re.match(r"^#\s+", lines[0]):
        lines = lines[1:]
    while lines and not lines[0].strip():
        lines.pop(0)
    return "\n".join(lines).strip()


def _chapter_heading(
    chapter: ExportedChapter,
    style: Literal["chinese", "arabic", "chapter", "plain"],
) -> str:
    if style == "arabic":
        return f"第{chapter.chapter_number}章 {chapter.title}"
    if style == "chapter":
        return f"Chapter {chapter.chapter_number} {chapter.title}"
    if style == "plain":
        return f"{chapter.chapter_number}. {chapter.title}"
    return f"第{_chapter_number_text(chapter.chapter_number)}章 {chapter.title}"


def _markdown_anchor(heading: str) -> str:
    cleaned = re.sub(r"[^\w\u4e00-\u9fff\s-]", "", heading, flags=re.UNICODE)
    return re.sub(r"\s+", "-", cleaned.strip().lower())


def _sha256_file(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as file:
        for chunk in iter(lambda: file.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest()


def _resolve_output_path(root: Path, output_path: Path | None, *, default_name: str) -> Path:
    if output_path is None:
        return root / "exports" / default_name
    if output_path.is_absolute():
        return output_path
    return root / output_path


def _rel(root: Path, path: Path) -> str:
    try:
        return str(path.relative_to(root))
    except ValueError:
        return str(path)
def _export_id(timestamp: datetime) -> str:
    return "export_" + timestamp.strftime("%Y%m%d_%H%M%S_%f")
